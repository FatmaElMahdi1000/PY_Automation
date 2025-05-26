import streamlit as st
from pathlib import Path
import PyPDF2
import docx
import pandas as pd
import io
import re

st.title("📄 Word Count Tool")
uploaded_files = st.file_uploader("Upload files", accept_multiple_files=True)

def clean_and_count_words(text):
    """Clean text and count words properly"""
    if not text or not text.strip():
        return 0
    
    # Remove extra whitespace and split by whitespace
    words = text.strip().split()
    # Filter out empty strings
    words = [word for word in words if word.strip()]
    return len(words)

def count_words(file, suffix):
    total_text = ""
    
    try:
        if suffix == ".txt":
            # Reset file pointer to beginning
            file.seek(0)
            total_text = file.read().decode("utf-8")
            
        elif suffix == ".pdf":
            # Reset file pointer to beginning
            file.seek(0)
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text = page.extract_text()
                if text: 
                    total_text += text + " "
                    
        elif suffix in [".docx", ".doc"]:
            # Reset file pointer to beginning
            file.seek(0)
            
            if suffix == ".docx":
                # Create a copy of the file content for python-docx
                file_content = io.BytesIO(file.read())
                doc_file = docx.Document(file_content)
                
                # Extract text from paragraphs
                for para in doc_file.paragraphs:
                    if para.text.strip():  # Only add non-empty paragraphs
                        total_text += para.text + " "
                
                # Also extract text from tables
                for table in doc_file.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                total_text += cell.text + " "
                                
                # Extract text from headers and footers
                for section in doc_file.sections:
                    if section.header:
                        for para in section.header.paragraphs:
                            if para.text.strip():
                                total_text += para.text + " "
                    if section.footer:
                        for para in section.footer.paragraphs:
                            if para.text.strip():
                                total_text += para.text + " "
            else:
                # For .doc files, python-docx might not work properly
                st.warning(f"⚠️ .doc files might not be fully supported. Consider converting to .docx format.")
                file_content = io.BytesIO(file.read())
                try:
                    doc_file = docx.Document(file_content)
                    for para in doc_file.paragraphs:
                        if para.text.strip():
                            total_text += para.text + " "
                except Exception as e:
                    st.error(f"Error reading .doc file: {str(e)}")
                    return 0
                    
        elif suffix in [".xlsx", ".csv"]:
            file.seek(0)
            if suffix == ".xlsx":
                df = pd.read_excel(file)
            else:
                df = pd.read_csv(file)
            
            # Convert all values to string and join
            total_text = " ".join(map(str, df.values.flatten()))
            
    except Exception as e:
        st.error(f"Error processing {file.name}: {str(e)}")
        return 0
    
    return clean_and_count_words(total_text)

if uploaded_files:
    total_wc = 0
    
    for file in uploaded_files:
        suffix = Path(file.name).suffix.lower()
        
        # Debug info
        st.write(f"Processing: **{file.name}** (Type: {suffix})")
        
        wc = count_words(file, suffix)
        st.write(f"**{file.name}** ➜ {wc} words")
        
        total_wc += wc
    
    st.success(f"✅ Total Word Count: {total_wc}")
    
    # Additional debug option
    if st.checkbox("Show debug information"):
        for file in uploaded_files:
            suffix = Path(file.name).suffix.lower()
            file.seek(0)
            
            if suffix in [".docx", ".doc"]:
                try:
                    file_content = io.BytesIO(file.read())
                    doc_file = docx.Document(file_content)
                    
                    st.write(f"**{file.name} Debug Info:**")
                    st.write(f"- Number of paragraphs: {len(doc_file.paragraphs)}")
                    st.write(f"- Number of tables: {len(doc_file.tables)}")
                    
                    # Show first few paragraphs
                    st.write("First few paragraphs:")
                    for i, para in enumerate(doc_file.paragraphs[:5]):
                        if para.text.strip():
                            st.write(f"  {i+1}: {para.text[:100]}...")
                            
                except Exception as e:
                    st.write(f"Debug error for {file.name}: {str(e)}")